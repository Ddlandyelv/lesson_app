"""docx 教案排版引擎 — 参照美术教案样本格式"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pathlib import Path
from typing import Optional

from .models import LessonData


class LessonTemplate:
    """参照现有美术教案样本排版"""

    def __init__(self):
        self.doc = Document()

        for section in self.doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)

        style = self.doc.styles["Normal"]
        style.font.name = "微软雅黑"
        style.font.size = Pt(12)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.paragraph_format.line_spacing = 1.5

    def _sf(self, run, size=12, bold=False):
        run.font.size = Pt(size)
        run.bold = bold
        run.font.name = "微软雅黑"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    def title(self, text):
        """标题：22pt 加粗 居中"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(20)
        r = p.add_run(text)
        self._sf(r, size=22, bold=True)

    def info(self, label, value):
        """信息行：12pt"""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run(label + value)
        self._sf(r, size=12)

    def heading(self, text):
        """小节标题（一、）：16pt 加粗"""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run(text)
        self._sf(r, size=16, bold=True)

    def body(self, text):
        """正文：12pt 首行缩进"""
        if not text.strip():
            return
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.75)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run(text.strip())
        self._sf(r, size=12)

    def label_body(self, text):
        """带标签正文（如【重点】）：12pt 首行缩进，不加粗"""
        if not text.strip():
            return
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.75)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run(text.strip())
        self._sf(r, size=12)

    def sub_heading(self, text):
        """子标题（如【步骤一】）：12pt 首行缩进，不加粗"""
        if not text.strip():
            return
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.75)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run(text.strip())
        self._sf(r, size=12)

    def bullet(self, text):
        """列表项：12pt 缩进"""
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run("• " + text)
        self._sf(r, size=12)

    def add_image(self, image_path, max_width=14):
        img_path = Path(image_path)
        if not img_path.exists():
            return
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run()
        r.add_picture(str(img_path), width=Cm(min(max_width, 14)))

    def render(self, data: LessonData, output_path: str, image_path: Optional[str] = None):
        self.title(f"《{data.title}》教案")

        self.info("适用班级：", data.level or "（由教师填写）")
        self.info("课    时：", data.duration or "（由教师填写）")
        self.info("绘画材料：", data.materials or "（由教师填写）")

        if image_path:
            self.add_image(image_path)

        # 教学目标
        self.heading("一、教学目标")
        for line in str(data.objectives).split("\n"):
            if line.strip():
                self.body(line)

        # 教学重难点
        self.heading("二、教学重难点")
        self.label_body(f"【教学重点】{data.key_points}")
        self.label_body(f"【教学难点】{data.difficult_points}")

        # 教学过程
        self.heading("三、教学过程")
        for step in str(data.process).split("\n"):
            step = step.strip()
            if not step:
                continue
            if step.startswith("【") and "】" in step[:10]:
                self.sub_heading(step)
            else:
                self.body(step)

        # 课后延伸
        if data.extension.strip():
            self.heading("四、课后延伸")
            self.body(data.extension)

        # 注意事项
        if data.notes.strip():
            self.heading("五、注意事项")
            for line in str(data.notes).split("\n"):
                if line.strip():
                    self.bullet(line)

        self.doc.save(output_path)
